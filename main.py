import json
import os
import re
import tempfile
import logging
import time
import asyncio
import datetime 
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel, Field
import PyPDF2
import docx
from typing import Dict, List, Optional, Literal, Any, Callable, Tuple
from pdf2image import convert_from_path
from PIL import Image

# Import metrics logging
from metrics_logger import RequestMetrics, log_model_metrics, log_file_processing_metrics, shutdown_metrics_logger

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
    handlers=[logging.StreamHandler()]
)

class UserFriendlyFilter(logging.Filter):
    def filter(self, record):
        if "Sending prompt to" in record.msg: record.msg = "🤖 Processing with Gemma 3:4b model..."
        elif "Received response from" in record.msg: record.msg = "✅ Received model response in {:.2f}s".format(float(record.msg.split("in ")[1].split("s")[0]))
        elif "Extracting text from PDF" in record.msg or "Extracting text from DOCX" in record.msg: record.msg = "📄 Extracting text from document..."
        elif "Successfully extracted" in record.msg and "characters from" in record.msg: record.msg = "✅ Successfully extracted document text"
        elif "Parsing resume with" in record.msg: record.msg = "🔍 Analyzing resume content..."
        elif "Parsing job description with" in record.msg: record.msg = "🔍 Analyzing job description content..."
        elif "Starting resume parsing" in record.msg: record.msg = "🚀 Starting resume analysis..."
        elif "Starting job description parsing" in record.msg: record.msg = "🚀 Starting job description analysis..."
        elif "Normalizing" in record.msg: record.msg = "📊 Organizing extracted data..."
        elif "Image fallback for" in record.msg: record.msg = f"↪️ {record.msg}" 
        return True

logger = logging.getLogger("gemma-api")
logger.addFilter(UserFriendlyFilter())

os.environ.pop('SSL_CERT_FILE', None)
import ollama
MODEL_NAME = "gemma3:4b"

MIN_TEXT_LENGTH_FOR_FALLBACK = 50 

# <editor-fold desc="Pydantic Models">
class EducationEntry(BaseModel):
    degree: str; institution: str; year: str; gpa: Optional[str] = None; location: Optional[str] = None; major: Optional[str] = None; minor: Optional[str] = None; achievements: Optional[List[str]] = None; courses: Optional[List[str]] = None
class ExperienceEntry(BaseModel):
    company_name: str; role: str; duration: str; key_responsibilities: str; location: Optional[str] = None; achievements: Optional[List[str]] = None; technologies_used: Optional[List[str]] = None; team_size: Optional[str] = None; industry: Optional[str] = None
class ProjectEntry(BaseModel):
    name: str; description: Optional[str] = None; duration: Optional[str] = None; technologies_used: Optional[List[str]] = None; url: Optional[str] = None; role: Optional[str] = None; achievements: Optional[List[str]] = None
class CertificationEntry(BaseModel):
    name: str; issuer: Optional[str] = None; date: Optional[str] = None; expiry: Optional[str] = None; url: Optional[str] = None
class SocialMediaEntry(BaseModel):
    platform: str; url: str; username: Optional[str] = None
class LanguageEntry(BaseModel):
    name: str; proficiency: Optional[str] = None
class PublicationEntry(BaseModel):
    title: str; publisher: Optional[str] = None; date: Optional[str] = None; url: Optional[str] = None; authors: Optional[List[str]] = None; description: Optional[str] = None
class AchievementEntry(BaseModel):
    title: str; date: Optional[str] = None; issuer: Optional[str] = None; description: Optional[str] = None
class VolunteerEntry(BaseModel):
    organization: str; role: str; duration: Optional[str] = None; description: Optional[str] = None

class ResumeResponse(BaseModel):
    name: str; email: Optional[str] = None; phone: Optional[str] = None; summary: Optional[str] = None; location: Optional[str] = None; education: List[EducationEntry] = []; skills: List[str] = []; direct_skills: Dict[str, str] = {}; subjective_skills: Dict[str, str] = {}; experience: List[ExperienceEntry] = []; projects: List[ProjectEntry] = []; certifications: List[CertificationEntry] = []; languages: List[LanguageEntry] = []; social_media: List[SocialMediaEntry] = []; publications: List[PublicationEntry] = []; achievements: List[AchievementEntry] = []; volunteer_experience: List[VolunteerEntry] = []; domain_of_interest: List[str] = []; references: List[Dict[str, str]] = []; confidence_score: float = 0.0; confidence_details: Dict[str, float] = {}
class BenefitEntry(BaseModel):
    title: str; description: Optional[str] = None
class RequirementEntry(BaseModel):
    title: str; description: Optional[str] = None; is_mandatory: Optional[bool] = True
class JobDescriptionResponse(BaseModel):
    job_title: str; company_name: Optional[str] = None; location: Optional[str] = None; job_type: Optional[str] = None; work_mode: Optional[str] = None; department: Optional[str] = None; summary: Optional[str] = None; responsibilities: List[str] = []; required_skills: List[str] = []; preferred_skills: List[str] = []; required_experience: Optional[str] = None; education_requirements: List[str] = []; salary_range: Optional[str] = None; benefits: List[BenefitEntry] = []; requirements: List[RequirementEntry] = []; application_deadline: Optional[str] = None; posting_date: Optional[str] = None; contact_information: Optional[Dict[str, str]] = None; company_description: Optional[str] = None; industry: Optional[str] = None; career_level: Optional[str] = None; confidence_score: float = 0.0
    class Config: extra = "allow"
class GenerateRequest(BaseModel):
    prompt: str; history: str = ""
class JDQuestionRequest(BaseModel):
    resume_json: Dict[str, Any] = Field(..., description="Resume data in JSON format"); technical_questions: int = Field(..., ge=0, le=10); past_experience_questions: int = Field(..., ge=0, le=10); case_study_questions: int = Field(..., ge=0, le=10); situation_handling_questions: int = Field(..., ge=0, le=10); personality_test_questions: int = Field(..., ge=0, le=10)
class InterfixRequest(BaseModel):
    summary: str = Field(..., description="Summary or transcript of VAPI response")
class InterfixResponse(BaseModel):
    offer_in_hand: Optional[float] = None; notice_period: Optional[str] = None; expected_salary: Optional[float] = None; reason_to_switch: Optional[str] = None; preferred_time_for_interview: Optional[str] = None; preferred_date_for_interview: Optional[str] = None
class IntervetRequest(BaseModel):
    resume_json: Dict[str, Any] = Field(..., description="Resume data"); jd_json: Dict[str, Any] = Field(..., description="Job description data")
# </editor-fold>

app = FastAPI(title="Gemma 3:4B API", description="API for interacting with Gemma 3:4B model and parsing resumes", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

class MetricsMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next: Callable) -> Response:
        metrics = RequestMetrics(endpoint=request.url.path)
        request.state.metrics = metrics
        try:
            response = await call_next(request)
            metrics.mark_complete(status_code=response.status_code)
            return response
        except Exception as e:
            metrics.mark_complete(status_code=500, error=str(e))
            raise
app.add_middleware(MetricsMiddleware)

from contextlib import asynccontextmanager # Ensure this is at the top-level imports if not already
@asynccontextmanager
async def lifespan(_: FastAPI): yield; shutdown_metrics_logger()
app.router.lifespan_context = lifespan

async def get_response(prompt: str, timeout_seconds: int = 60, max_tokens: int = 1000, image_path: str = None, request_metrics: RequestMetrics = None) -> str:
    try:
        logger.info(f"Sending prompt to {MODEL_NAME} with {timeout_seconds}s timeout")
        start_time = time.time()
        options = {"num_predict": max_tokens}

        def ollama_call_logic():
            if image_path and os.path.exists(image_path):
                logger.info(f"Including image from path: {image_path}")
                try:
                    return ollama.generate(model=MODEL_NAME, prompt=prompt, images=[image_path], stream=False, options=options)
                except Exception as img_error:
                    logger.warning(f"Error using images parameter: {img_error}. Trying alternative method...")
                    with open(image_path, 'rb') as img_file:
                        import base64
                        img_data = base64.b64encode(img_file.read()).decode('utf-8')
                        special_prompt = f"[img]{img_data}[/img]\n{prompt}"
                        return ollama.generate(model=MODEL_NAME, prompt=special_prompt, stream=False, options=options)
            else:
                return ollama.generate(model=MODEL_NAME, prompt=prompt, stream=False, options=options)
        try:
            result = await asyncio.wait_for(asyncio.to_thread(ollama_call_logic), timeout=timeout_seconds)
            processing_time = time.time() - start_time
            logger.info(f"Received response from {MODEL_NAME} in {processing_time:.2f}s")
            response_text = result["response"]
            if request_metrics:
                request_metrics.mark_first_byte()
                log_model_metrics(request_metrics=request_metrics, model_name=MODEL_NAME, prompt_length=len(prompt), response_length=len(response_text), processing_time=processing_time, error=None)
                if "eval_count" in result: request_metrics.add_metric("eval_count", result["eval_count"])
                if "prompt_eval_count" in result: request_metrics.add_metric("prompt_eval_count", result["prompt_eval_count"])
                if "total_duration" in result: request_metrics.add_metric("model_total_duration", result["total_duration"])
            return response_text
        except asyncio.TimeoutError:
            logger.warning(f"Response generation timed out after {timeout_seconds}s for prompt: {prompt[:100]}...")
            if request_metrics: log_model_metrics(request_metrics=request_metrics, model_name=MODEL_NAME, prompt_length=len(prompt), response_length=0, processing_time=timeout_seconds, error=f"Timeout after {timeout_seconds}s")
            raise HTTPException(status_code=504, detail=f"Response generation timed out after {timeout_seconds} seconds.")
        except Exception as e:
            logger.error(f"Error in model generation for prompt {prompt[:100]}...: {e}", exc_info=True)
            if request_metrics: log_model_metrics(request_metrics=request_metrics, model_name=MODEL_NAME, prompt_length=len(prompt), response_length=0, processing_time=(time.time() - start_time), error=str(e))
            raise HTTPException(status_code=500, detail=f"Error generating response: {str(e)}")
    except HTTPException: raise
    except Exception as e:
        logger.error(f"Error in get_response setup: {e}", exc_info=True)
        if request_metrics: log_model_metrics(request_metrics=request_metrics, model_name=MODEL_NAME, prompt_length=len(prompt) if 'prompt' in locals() else 0, response_length=0, processing_time=time.time() - start_time if 'start_time' in locals() else 0, error=str(e))
        raise HTTPException(status_code=500, detail=f"Internal server error in get_response: {str(e)}")

async def extract_text_from_pdf_standard(file_path: str) -> Optional[str]:
    try:
        logger.info(f"Extracting text from PDF (standard): {file_path}")
        def _extract_sync():
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages: text += (page.extract_text() or "") + "\n"
            return text
        text = await asyncio.to_thread(_extract_sync)
        if not text.strip() or len(text.strip()) < MIN_TEXT_LENGTH_FOR_FALLBACK: 
            logger.info(f"Standard PDF extraction for {file_path} yielded insufficient text (length: {len(text.strip())}).")
            return None 
        logger.info(f"Successfully extracted {len(text)} characters from PDF (standard): {file_path}")
        return text
    except Exception as e: logger.error(f"Error extracting text from PDF (standard) for {file_path}: {e}", exc_info=True); return None

async def extract_text_from_docx_standard(file_path: str) -> Optional[str]:
    try:
        logger.info(f"Extracting text from DOCX (standard): {file_path}")
        def _extract_sync():
            doc = docx.Document(file_path)
            local_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    local_text += "\n" + " | ".join([cell.text for cell in row.cells if cell.text.strip()])
            return local_text
        text = await asyncio.to_thread(_extract_sync)
        if not text.strip() or len(text.strip()) < MIN_TEXT_LENGTH_FOR_FALLBACK: 
            logger.info(f"Standard DOCX extraction for {file_path} yielded insufficient text (length: {len(text.strip())}).")
            return None
        logger.info(f"Successfully extracted {len(text)} characters from DOCX (standard): {file_path}")
        return text
    except Exception as e: logger.error(f"Error extracting text from DOCX (standard) for {file_path}: {e}", exc_info=True); return None

async def extract_text_from_image(file_path: str, parsing_prompt_override: Optional[str] = None) -> str:
    logger.info(f"Attempting to extract text from file as image: {file_path}")
    try:
        current_prompt = parsing_prompt_override
        if not current_prompt:
            current_prompt = "This document contains important text. Please extract ALL text content from this image..." 
            logger.info("Using generic image text extraction prompt.")
        else:
            logger.info("Using parsing_prompt_override for image text extraction.")

        response = await get_response(prompt=current_prompt, timeout_seconds=120, max_tokens=2000, image_path=file_path)
        text = response.strip()
        if len(text) < MIN_TEXT_LENGTH_FOR_FALLBACK: 
            logger.warning(f"Image-based text extraction for {file_path} returned insufficient content (length: {len(text)}).")
            raise ValueError("Image-based text extraction returned insufficient content")
        logger.info(f"Successfully extracted {len(text)} characters using image-based extraction for {file_path} with provided prompt.")
        return text
    except Exception as e: 
        logger.error(f"Error in image-based text extraction for {file_path}: {e}", exc_info=True)
        raise ValueError(f"Failed to extract text using image-based approach from {file_path}: {str(e)}")


async def extract_text_from_document_via_image_fallback(
    file_path: str, 
    file_type: Literal["pdf", "docx"], 
    parsing_prompt: str, 
    request_metrics: Optional[RequestMetrics] = None
) -> Optional[str]:
    logger.info(f"Image fallback for {file_type}: {file_path}")
    full_text_from_images = ""

    if file_type == "pdf":
        try:
            images = await asyncio.to_thread(convert_from_path, file_path)
            logger.info(f"Converted PDF {file_path} to {len(images)} images for fallback processing.")
            for i, image_pil in enumerate(images):
                page_image_path = None
                try:
                    def save_pil_image(): 
                        _page_image_path = tempfile.mktemp(suffix=".jpg")
                        image_pil.save(_page_image_path, "JPEG")
                        return _page_image_path
                    page_image_path = await asyncio.to_thread(save_pil_image)
                    logger.info(f"Extracting text from page image {i+1}/{len(images)} for {file_path}")
                    page_text = await extract_text_from_image(page_image_path, parsing_prompt_override=parsing_prompt)
                    full_text_from_images += page_text + "\n\n"
                except Exception as page_err:
                    logger.warning(f"Could not extract text from page image {i+1} of {file_path}: {page_err}", exc_info=True)
                finally:
                    if page_image_path and await asyncio.to_thread(os.path.exists, page_image_path):
                        await asyncio.to_thread(os.remove, page_image_path)
            if request_metrics: request_metrics.add_metric("image_fallback_pages_processed", len(images))
        except Exception as e:
            logger.error(f"PDF to image conversion failed for {file_path}: {e}", exc_info=True)
            if request_metrics: request_metrics.add_metric("image_fallback_error", f"PDF to image conversion failed for {file_path}: {e}")
            return None
    
    elif file_type == "docx":
        logger.warning(f"DOCX image fallback is simplified for {file_path}: attempting to treat whole DOCX as one image.")
        try:
            full_text_from_images = await extract_text_from_image(file_path, parsing_prompt_override=parsing_prompt)
            if request_metrics: request_metrics.add_metric("image_fallback_docx_attempted", True)
        except Exception as e:
            logger.error(f"DOCX direct image fallback failed for {file_path}: {e}", exc_info=True)
            if request_metrics: request_metrics.add_metric("image_fallback_error", f"DOCX direct image fallback failed for {file_path}: {e}")
            return None

    if not full_text_from_images.strip() or len(full_text_from_images.strip()) < MIN_TEXT_LENGTH_FOR_FALLBACK:
        logger.warning(f"Image fallback for {file_path} did not yield sufficient text.")
        return None
        
    logger.info(f"Successfully extracted {len(full_text_from_images)} characters via image fallback for {file_path}.")
    return full_text_from_images


async def extract_text_from_file(
    file_path: str, 
    file_type: Literal["pdf", "docx"], 
    parsing_prompt_override: str, 
    request_metrics: RequestMetrics = None
) -> str:
    start_time = time.time()
    file_size = await asyncio.to_thread(os.path.getsize, file_path) if await asyncio.to_thread(os.path.exists, file_path) else 0
    
    text = None
    extraction_method = "failed"
    standard_error_detail = "N/A"

    try:
        if file_type == "pdf":
            text = await extract_text_from_pdf_standard(file_path)
            extraction_method = "pdf_standard"
        elif file_type == "docx":
            text = await extract_text_from_docx_standard(file_path)
            extraction_method = "docx_standard"
        else:
            logger.error(f"Unsupported file type passed to extract_text_from_file: {file_type} for {file_path}")
            raise ValueError(f"Unsupported file type: {file_type}")

        if text: 
            logger.info(f"Standard extraction successful for {file_type}: {file_path}")
            if request_metrics: log_file_processing_metrics(request_metrics, file_type, file_size, extraction_method, len(text), time.time() - start_time)
            return text
        else:
            logger.info(f"Standard {file_type} extraction yielded insufficient text for {file_path}. Attempting image fallback.")
            if request_metrics: request_metrics.add_metric("standard_extraction_insufficient", True)
            standard_error_detail = "Standard extraction yielded insufficient text."
    
    except Exception as standard_err: 
        logger.warning(f"Standard {file_type} extraction failed for {file_path}: {standard_err}", exc_info=True)
        if request_metrics: request_metrics.add_metric("standard_extraction_error", str(standard_err))
        standard_error_detail = str(standard_err)

    try:
        text = await extract_text_from_document_via_image_fallback(file_path, file_type, parsing_prompt_override, request_metrics)
        if text:
            logger.info(f"Image fallback extraction successful for {file_type}: {file_path}.")
            extraction_method = "image_fallback"
            if request_metrics: log_file_processing_metrics(request_metrics, file_type, file_size, extraction_method, len(text), time.time() - start_time)
            return text
        else:
            logger.error(f"Image fallback also failed or yielded insufficient text for {file_type}: {file_path}.")
            raise HTTPException(status_code=400, detail=f"Could not extract sufficient text from {os.path.basename(file_path)} using any method. Standard error: {standard_error_detail}, Fallback: Insufficient text.")

    except Exception as fallback_err: 
        logger.error(f"Image fallback extraction for {file_type} {file_path} failed: {fallback_err}", exc_info=True)
        if request_metrics: log_file_processing_metrics(request_metrics, file_type, file_size, "failed_fallback", 0, time.time() - start_time); request_metrics.add_metric("image_fallback_final_error", str(fallback_err))
        raise HTTPException(status_code=400, detail=f"Text extraction failed for {os.path.basename(file_path)}. Standard error: {standard_error_detail}, Fallback error: {fallback_err}")


async def calculate_resume_confidence_score(parsed_data: Dict) -> float:
    field_weights = {"name": 0.10, "email": 0.10, "phone": 0.08, "summary": 0.05, "location": 0.07, "education": 0.12, "skills": 0.15, "experience": 0.15, "projects": 0.08, "certifications": 0.05, "languages": 0.05}
    score = 0.0; total_weight = 0.0; detailed_scores = {}
    for field, weight in field_weights.items():
        total_weight += weight; field_score = 0.0
        if field not in parsed_data or parsed_data[field] is None: detailed_scores[field] = 0.0; continue
        if parsed_data[field]: field_score = 0.8 
        score += weight * field_score; detailed_scores[field] = round(field_score * 100, 1)
    return round(score / total_weight if total_weight > 0 else 0, 2)

async def calculate_jd_confidence_score(parsed_data: Dict) -> float:
    field_weights = {"job_title": 0.12, "company_name": 0.10, "location": 0.08, "summary": 0.08, "responsibilities": 0.15, "required_skills": 0.15, "preferred_skills": 0.10, "required_experience": 0.07, "education_requirements": 0.05}
    score = 0.0; total_weight = 0.0
    for field, weight in field_weights.items():
        total_weight += weight; field_score = 0.0
        if parsed_data.get(field): field_score = 0.8
        score += weight * field_score
    return round(score / total_weight if total_weight > 0 else 0, 2)

async def convert_skills_to_dict(parsed_data: Dict) -> Dict:
    if "skills" not in parsed_data or not isinstance(parsed_data["skills"], list): return parsed_data
    skills_list = parsed_data["skills"]; direct_skills = {}; subjective_skills = {}
    for skill in skills_list:
        if isinstance(skill, str): direct_skills[skill] = "Mentioned in resume"
    parsed_data["direct_skills"] = direct_skills; parsed_data["subjective_skills"] = subjective_skills
    return parsed_data

RESUME_PARSING_PROMPT_TEMPLATE = """ You are an expert resume parser. Your task is to extract ALL structured information from the resume text below... Respond ONLY with the JSON object...""" 
JD_PARSING_PROMPT_TEMPLATE = """ You are an expert job description parser. Your task is to extract ALL structured information from the job description text below... Respond ONLY with the JSON object...""" 

async def parse_resume_with_gemma(resume_text_for_llm: str) -> Dict:
    logger.info("Parsing resume content with Gemma model")
    prompt = RESUME_PARSING_PROMPT_TEMPLATE.replace("... Resume text...", resume_text_for_llm) 
    try:
        response_text = await get_response(prompt, timeout_seconds=120)
        json_str = response_text.strip()
        if "```json" in json_str: json_str = json_str.split("```json")[1].split("```")[0].strip()
        elif "```" in json_str: json_str = json_str.split("```")[1].split("```")[0].strip()
        if not (json_str.startswith('{') and json_str.endswith('}')): raise json.JSONDecodeError("No JSON object found", json_str, 0)
        return await asyncio.to_thread(json.loads, json_str)
    except json.JSONDecodeError as json_err:
        logger.error(f"JSON decode error in parse_resume_with_gemma: {json_err.msg} for response: {response_text[:200] if 'response_text' in locals() else 'N/A'}", exc_info=True)
        return {"error": "Failed to parse resume JSON from LLM", "details": str(json_err)}
    except Exception as e:
        logger.error(f"Error in Gemma resume parsing: {e}", exc_info=True)
        return {"error": "Failed to parse resume with LLM", "details": str(e)}

async def parse_jd_with_gemma(jd_text_for_llm: str) -> Dict:
    logger.info("Parsing JD content with Gemma model")
    prompt = JD_PARSING_PROMPT_TEMPLATE.replace("... Job Description text...", jd_text_for_llm)
    try:
        response_text = await get_response(prompt, timeout_seconds=120)
        json_str = response_text.strip()
        if "```json" in json_str: json_str = json_str.split("```json")[1].split("```")[0].strip()
        elif "```" in json_str: json_str = json_str.split("```")[1].split("```")[0].strip()
        if not (json_str.startswith('{') and json_str.endswith('}')): raise json.JSONDecodeError("No JSON object found", json_str, 0)
        return await asyncio.to_thread(json.loads, json_str)
    except json.JSONDecodeError as json_err:
        logger.error(f"JSON decode error in parse_jd_with_gemma: {json_err.msg} for response: {response_text[:200] if 'response_text' in locals() else 'N/A'}", exc_info=True)
        return {"error": "Failed to parse JD JSON from LLM", "details": str(json_err)}
    except Exception as e:
        logger.error(f"Error in Gemma JD parsing: {e}", exc_info=True)
        return {"error": "Failed to parse JD with LLM", "details": str(e)}

async def normalize_resume_data(parsed_data: Dict, convert_skills_to_dict_format: bool = False) -> Dict:
    logger.info("Normalizing resume data")
    if not isinstance(parsed_data, dict): parsed_data = {"name":"Unknown", "error": "Invalid input to normalize_resume_data", "details": str(parsed_data)}
    if 'name' not in parsed_data or parsed_data['name'] is None:
        parsed_data['name'] = "Unknown"
        if 'error' not in parsed_data: parsed_data['error'] = "Candidate name could not be determined."
    normalized_dict = {key: parsed_data.get(key, ResumeResponse.__fields__[key].default) for key in ResumeResponse.__fields__}
    if convert_skills_to_dict_format:
         normalized_dict = await convert_skills_to_dict(normalized_dict)
    normalized_dict["confidence_score"] = await calculate_resume_confidence_score(normalized_dict)
    normalized_dict["confidence_details"] = {} 
    try:
        return ResumeResponse(**normalized_dict).dict(exclude_none=True)
    except Exception as validation_error:
        logger.error(f"Pydantic validation error during resume normalization: {validation_error}", exc_info=True)
        minimal_error_response = {"name": normalized_dict.get("name","Unknown (validation error)")}
        for field in ResumeResponse.__fields__:
            if field not in minimal_error_response:
                minimal_error_response[field] = ResumeResponse.__fields__[field].default
        minimal_error_response["error"] = "Data normalization/validation failed" 
        return ResumeResponse(**minimal_error_response).dict(exclude_none=True)


async def normalize_jd_data(parsed_data: Dict) -> Dict:
    logger.info("Normalizing JD data")
    if not isinstance(parsed_data, dict): parsed_data = {"job_title":"Unknown", "error": "Invalid input to normalize_jd_data", "details": str(parsed_data)}
    if 'job_title' not in parsed_data or parsed_data['job_title'] is None:
        parsed_data['job_title'] = "Unknown"
        if 'error' not in parsed_data:
            parsed_data['error'] = "Job title could not be determined."
    normalized_dict = {key: parsed_data.get(key, JobDescriptionResponse.__fields__[key].default) for key in JobDescriptionResponse.__fields__}
    normalized_dict["confidence_score"] = await calculate_jd_confidence_score(normalized_dict)
    normalized_dict["confidence_details"] = {}
    try:
        return JobDescriptionResponse(**normalized_dict).dict(exclude_none=True)
    except Exception as validation_error:
        logger.error(f"Pydantic validation error during JD normalization: {validation_error}", exc_info=True)
        minimal_error_response = {"job_title": normalized_dict.get("job_title","Unknown (validation error)")}
        for field in JobDescriptionResponse.__fields__:
             if field not in minimal_error_response:
                minimal_error_response[field] = JobDescriptionResponse.__fields__[field].default
        minimal_error_response["error"] = "Data normalization/validation failed"
        return JobDescriptionResponse(**minimal_error_response).dict(exclude_none=True)


async def parse_resume(file_path: str, file_type: Literal["pdf", "docx"], request_metrics: Optional[RequestMetrics] = None) -> Dict:
    logger.info(f"Starting resume parsing orchestration for {file_path}")
    specific_resume_parsing_prompt = RESUME_PARSING_PROMPT_TEMPLATE.replace("... Resume text...", "{text_content_placeholder}")
    try:
        extracted_text = await extract_text_from_file(file_path, file_type, specific_resume_parsing_prompt, request_metrics)
        if not extracted_text or extracted_text.strip() == "":
             logger.error(f"No text could be extracted from resume {os.path.basename(file_path)} by any method.")
             return await normalize_resume_data({"name": "Unknown", "error": f"No text extracted from resume: {os.path.basename(file_path)}"}, True)
        parsed_result = await parse_resume_with_gemma(extracted_text) 
        if "error" in parsed_result: logger.warning(f"Resume parsing with Gemma for {os.path.basename(file_path)} returned error: {parsed_result.get('error')}")
        return await normalize_resume_data(parsed_result, True)
    except HTTPException as http_exc:
        logger.error(f"HTTPException during resume processing for {os.path.basename(file_path)}: {http_exc.detail}", exc_info=True)
        return await normalize_resume_data({"name": "Unknown", "error": f"Resume processing error: {http_exc.detail}"}, True)
    except Exception as e:
        logger.error(f"Error in resume parsing pipeline for {os.path.basename(file_path)}: {e}", exc_info=True)
        return await normalize_resume_data({"name": "Unknown", "error": f"Unexpected error processing resume: {str(e)}"}, True)

async def extract_skills_from_text(text: str) -> List[str]: 
    def _extract_sync():
        logger.info("Attempting fallback skills extraction from text (sync)")
        common_skills = ["Python", "Java", "JavaScript", "C++", "C#", "SQL", "AWS", "Docker", "Kubernetes", "Machine Learning", "Deep Learning", "React", "Angular", "Node.js"]
        extracted_skills = [skill for skill in common_skills if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text.lower())]
        logger.info(f"Fallback extraction found {len(extracted_skills)} skills")
        return extracted_skills
    return await asyncio.to_thread(_extract_sync)

async def extract_education_from_text(text: str) -> List[str]: 
    def _extract_sync():
        logger.info("Attempting fallback education extraction from text (sync)")
        matches = re.findall(r"(?:bachelor|master|phd|b\.s|b\.a|m\.s|m\.a|b\.tech|m\.tech|b\.e|m\.e|associate)['']?s?\s*(?:degree)?\s*(?:in\s+[\w\s]+)?", text.lower(), re.IGNORECASE)
        formatted_edu = list(set([m.strip().capitalize() for m in matches]))
        logger.info(f"Fallback extraction found {len(formatted_edu)} education requirements")
        return formatted_edu
    return await asyncio.to_thread(_extract_sync)
    
async def parse_jd(file_path: str, file_type: Literal["pdf", "docx"], request_metrics: Optional[RequestMetrics] = None) -> Dict:
    logger.info(f"Starting JD parsing orchestration for {file_path}")
    specific_jd_parsing_prompt = JD_PARSING_PROMPT_TEMPLATE.replace("... Job Description text...", "{text_content_placeholder}")
    try:
        extracted_text = await extract_text_from_file(file_path, file_type, specific_jd_parsing_prompt, request_metrics)
        if not extracted_text or extracted_text.strip() == "":
            logger.error(f"No text could be extracted from JD {os.path.basename(file_path)} by any method.")
            return await normalize_jd_data({"job_title": "Unknown", "error": f"No text extracted from JD: {os.path.basename(file_path)}"})
        parsed_result = await parse_jd_with_gemma(extracted_text)
        if "error" in parsed_result: logger.warning(f"JD parsing with Gemma for {os.path.basename(file_path)} returned error: {parsed_result.get('error')}")
        normalized_result = await normalize_jd_data(parsed_result)
        if not normalized_result.get("required_skills") and isinstance(extracted_text, str):
            normalized_result["required_skills"] = await extract_skills_from_text(extracted_text)
        if not normalized_result.get("education_requirements") and not normalized_result.get("education_details",{}).get("degree_level") and isinstance(extracted_text, str):
            normalized_result["education_requirements"] = await extract_education_from_text(extracted_text)
        return normalized_result
    except HTTPException as http_exc:
        logger.error(f"HTTPException during JD processing for {os.path.basename(file_path)}: {http_exc.detail}", exc_info=True)
        return await normalize_jd_data({"job_title": "Unknown", "error": f"JD processing error: {http_exc.detail}"})
    except Exception as e:
        logger.error(f"Error in JD parsing pipeline for {os.path.basename(file_path)}: {e}", exc_info=True)
        return await normalize_jd_data({"job_title": "Unknown", "error": f"Unexpected error processing JD: {str(e)}"})

async def write_to_temp_file(content: bytes, suffix: str) -> str:
    def _write_sync():
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_f:
            temp_f.write(content)
            return temp_f.name
    return await asyncio.to_thread(_write_sync)

@app.get("/")
async def root(): return {"message": "Welcome to the Gemma3 API", "endpoints": []}

@app.post("/generate")
async def generate_endpoint(request_body: GenerateRequest):
    try:
        full_prompt = (request_body.history or "") + "\nUser: " + request_body.prompt + "\nAssistant: "
        response = await get_response(full_prompt)
        return {"response": response}
    except HTTPException as e: raise e
    except Exception as e: logger.error(f"Unexpected error in /generate: {e}", exc_info=True); raise HTTPException(status_code=500, detail="Unexpected error.")

@app.post("/generate_with_image")
async def generate_with_image_endpoint(file: UploadFile = File(...), prompt: str = Form(...), history: str = Form("")):
    temp_file_path = None
    try:
        file_extension = os.path.splitext(file.filename.lower())[1]
        supported_extensions = [".jpg", ".jpeg", ".png", ".gif", ".bmp"]
        if file_extension not in supported_extensions: raise HTTPException(status_code=400, detail="Unsupported image format")
        content = await file.read()
        if not content: raise HTTPException(status_code=400, detail="Empty file uploaded.")
        temp_file_path = await write_to_temp_file(content, file_extension)
        image_prompt_for_llm = f"I'm looking at this image. {prompt}" 
        if history: image_prompt_for_llm = f"{history}\n\n{image_prompt_for_llm}"
        response = await get_response(prompt=image_prompt_for_llm, image_path=temp_file_path, timeout_seconds=120, max_tokens=1500)
        return {"response": response}
    except HTTPException as e: raise e
    except Exception as e: logger.error(f"Unexpected error in /generate_with_image for {file.filename if file else 'N/A'}: {e}", exc_info=True); raise HTTPException(status_code=500, detail=str(e))
    finally:
        if temp_file_path and await asyncio.to_thread(os.path.exists, temp_file_path): await asyncio.to_thread(os.remove, temp_file_path)

@app.post("/resume", response_model=Dict)
async def parse_resume_endpoint(request: Request, file: UploadFile = File(...)):
    metrics = getattr(request.state, "metrics", None)
    temp_file_path = None
    try:
        file_extension = os.path.splitext(file.filename.lower())[1]
        if file_extension == '.pdf': file_type = "pdf"; suffix = '.pdf'
        elif file_extension == '.docx': file_type = "docx"; suffix = '.docx'
        else: raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
        content = await file.read()
        temp_file_path = await write_to_temp_file(content, suffix)
        parsed_data = await parse_resume(temp_file_path, file_type, request_metrics=metrics)
        return parsed_data
    except HTTPException as e: 
        if metrics: metrics.add_metric("error", str(e.detail if hasattr(e, 'detail') else e)); metrics.add_metric("error_status", e.status_code if hasattr(e, 'status_code') else 500)
        raise e
    except Exception as e: 
        if metrics: metrics.add_metric("error", str(e)); metrics.add_metric("error_status", 500)
        logger.error(f"Error in /resume for {file.filename}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred while processing {file.filename}: {str(e)}")
    finally:
        if temp_file_path and await asyncio.to_thread(os.path.exists, temp_file_path): await asyncio.to_thread(os.remove, temp_file_path)

async def generate_questions_for_category(jd_text: str, resume_data: Optional[Dict], category: str, num_questions: int) -> List[str]:
    if num_questions <= 0: return []
    logger.info(f"Generating {num_questions} questions for {category}")
    category_name_map = {"technical_questions": "Technical", "past_experience_questions": "Past Experience", "case_study_questions": "Case Study", "situation_handling_questions": "Situation Handling", "personality_test_questions": "Personality Test"}
    category_name = category_name_map.get(category, category)
    prompt = f"You are an expert interview question generator. Your task is to create {num_questions} {category_name} questions."
    if category == "personality_test_questions":
        prompt += " Focus on general personality traits. Candidate Resume Data (if available):\n" + (json.dumps(resume_data, indent=2) if resume_data else "N/A")
    else:
        prompt += f" Based on Job Description:\n{jd_text}\nAnd Candidate Resume Data (if available):\n" + (json.dumps(resume_data, indent=2) if resume_data else "N/A")
    prompt += "\nIMPORTANT: Respond ONLY with a JSON array of strings. Example: [\"Question 1?\"]"
    try:
        response_text = await get_response(prompt, timeout_seconds=45, max_tokens=500)
        json_str = response_text.strip()
        if "```json" in json_str: json_str = json_str.split("```json")[1].split("```")[0].strip()
        elif "```" in json_str: json_str = json_str.split("```")[1].split("```")[0].strip()
        if not (json_str.startswith('[') and json_str.endswith(']')): logger.warning(f"Response for {category} not a valid JSON array: {json_str[:100]}..."); return []
        questions = await asyncio.to_thread(json.loads, json_str)
        return [str(q) for q in questions if q][:num_questions] if isinstance(questions, list) else []
    except Exception as e: logger.error(f"Error generating questions for {category}: {e}", exc_info=True); return []

async def generate_interview_questions(jd_text: str, resume_data: Dict, question_scales: Dict) -> Dict:
    logger.info("Starting interview question generation (async)")
    categories = ["technical_questions", "past_experience_questions", "case_study_questions", "situation_handling_questions", "personality_test_questions"]
    questions_data = {}
    async def process_category(cat):
        scale = question_scales.get(cat, 0); num_q = min(5, max(0, int(scale / 2)))
        return cat, await generate_questions_for_category(jd_text, resume_data, cat, num_q)
    results = await asyncio.gather(*(process_category(cat) for cat in categories))
    for category, questions in results: questions_data[category] = questions
    logger.info("Successfully generated all interview questions")
    return questions_data

@app.post("/jd_parser", response_model=Dict)
async def parse_jd_endpoint(request: Request, file: UploadFile = File(...)):
    metrics = getattr(request.state, "metrics", None)
    temp_file_path = None
    try:
        file_extension = os.path.splitext(file.filename.lower())[1]
        if file_extension == '.pdf': file_type = "pdf"; suffix = '.pdf'
        elif file_extension == '.docx': file_type = "docx"; suffix = '.docx'
        else: raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
        content = await file.read()
        temp_file_path = await write_to_temp_file(content, suffix)
        parsed_data = await parse_jd(temp_file_path, file_type, request_metrics=metrics)
        return parsed_data
    except HTTPException as e: 
        if metrics: metrics.add_metric("error", str(e.detail if hasattr(e, 'detail') else e)); metrics.add_metric("error_status", e.status_code if hasattr(e, 'status_code') else 500)
        raise e
    except Exception as e: 
        if metrics: metrics.add_metric("error", str(e)); metrics.add_metric("error_status", 500)
        logger.error(f"Error in /jd_parser for {file.filename}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred while processing {file.filename}: {str(e)}")
    finally:
        if temp_file_path and await asyncio.to_thread(os.path.exists, temp_file_path): await asyncio.to_thread(os.remove, temp_file_path)

@app.post("/jd", response_model=Dict)
async def generate_interview_questions_endpoint(file: UploadFile = File(...), request_data_json: str = Form(...)):
    temp_file_path = None
    try:
        request_data_dict = await asyncio.to_thread(json.loads, request_data_json)
        request_data_obj = JDQuestionRequest(**request_data_dict)
        file_extension = os.path.splitext(file.filename.lower())[1]
        if file_extension == '.pdf': file_type = "pdf"; suffix = '.pdf'
        elif file_extension == '.docx': file_type = "docx"; suffix = '.docx'
        else: raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported for JD.")
        content = await file.read()
        temp_file_path = await write_to_temp_file(content, suffix)
        generic_text_extraction_prompt = "Extract all text content from this document."
        jd_text = await extract_text_from_file(temp_file_path, file_type, generic_text_extraction_prompt)
        question_scales = {cat: getattr(request_data_obj, cat) for cat in JDQuestionRequest.__fields__ if cat.endswith("_questions")}
        questions_data = await generate_interview_questions(jd_text, request_data_obj.resume_json, question_scales)
        return questions_data
    except HTTPException as e: raise e
    except Exception as e: logger.error(f"Error in /jd for {file.filename if file else 'N/A'}: {e}", exc_info=True); raise HTTPException(status_code=500, detail=str(e))
    finally:
        if temp_file_path and await asyncio.to_thread(os.path.exists, temp_file_path): await asyncio.to_thread(os.remove, temp_file_path)

# --- Candidate-Job Fit Evaluation Refactor ---
def normalize_skill(skill: str) -> str:
    if not skill: return ""
    normalized = str(skill).lower()
    normalized = re.sub(r'\b(programming|development|developer|engineer|engineering|specialist|expert|proficiency|basics of|basic|advanced|intermediate|experience with|experience in|knowledge of|skills in)\b', '', normalized)
    normalized = re.sub(r'[^\w\s]', '', normalized)
    return re.sub(r'\s+', ' ', normalized).strip()

def get_skill_variations(skill: str) -> List[str]:
    if not skill: return []
    variations = [skill]; normalized = normalize_skill(skill)
    if not normalized: return variations
    if normalized != str(skill).lower(): variations.append(normalized)
    if normalized in ["python", "java"]: variations.extend([f"{normalized} programming"])
    return list(set(variations))

def is_skill_match(resume_skill: str, jd_skill: str) -> bool:
    if not resume_skill or not jd_skill: return False
    resume_skill_norm = normalize_skill(resume_skill); jd_skill_norm = normalize_skill(jd_skill)
    if not resume_skill_norm or not jd_skill_norm: return False
    if resume_skill_norm == jd_skill_norm: return True
    if len(resume_skill_norm) > 2 and len(jd_skill_norm) > 2:
        if (resume_skill_norm in jd_skill_norm) or (jd_skill_norm in resume_skill_norm): return True
    resume_variations = get_skill_variations(resume_skill); jd_variations = get_skill_variations(jd_skill)
    for rv in resume_variations:
        if not rv: continue
        for jv in jd_variations:
            if not jv: continue
            if rv.lower() == jv.lower(): return True
            if len(rv) > 2 and len(jv) > 2 and (rv.lower() in jv.lower() or jv.lower() in rv.lower()): return True
    return False

def find_matching_skills(resume_skills: List[str], jd_skills: List[str]) -> Tuple[List[str], List[str]]:
    if resume_skills is None: resume_skills = []
    if jd_skills is None: jd_skills = []
    matched = []; unmatched = []
    valid_resume_skills = [s for s in resume_skills if s]; valid_jd_skills = [s for s in jd_skills if s]
    for jd_skill in valid_jd_skills:
        found_match = any(is_skill_match(resume_skill, jd_skill) for resume_skill in valid_resume_skills)
        if found_match: matched.append(jd_skill)
        else: unmatched.append(jd_skill)
    return matched, unmatched

def _prepare_llm_evaluation_context(resume_data: Dict, jd_data: Dict) -> Dict:
    logger.info("Preparing LLM evaluation context...")
    context = {}
    resume_direct_skills_list = list(resume_data.get("direct_skills", {}).keys())
    resume_subjective_skills_list = list(resume_data.get("subjective_skills", {}).keys())
    jd_required_skills = jd_data.get("required_skills", []); jd_preferred_skills = jd_data.get("preferred_skills", [])
    all_jd_skills = jd_required_skills + jd_preferred_skills
    matched_required, missing_required = find_matching_skills(resume_direct_skills_list, jd_required_skills)
    matched_preferred, _ = find_matching_skills(resume_direct_skills_list, jd_preferred_skills)
    subjective_matches_in_jd = [s for s in resume_subjective_skills_list if any(is_skill_match(s, jd_s) for jd_s in all_jd_skills) and not any(is_skill_match(s, direct_s) for direct_s in resume_direct_skills_list)]
    context["matched_required_skills"] = matched_required; context["missing_required_skills"] = missing_required
    context["matched_preferred_skills"] = matched_preferred; context["subjective_skills_found_in_jd"] = list(set(subjective_matches_in_jd))
    candidate_yoe = 0
    if isinstance(resume_data.get("experience"), list):
        for exp in resume_data["experience"]:
            if isinstance(exp, dict) and "duration" in exp:
                duration = exp["duration"]
                if isinstance(duration, str):
                    year_match = re.search(r'(\d{4})\s*-\s*(\d{4}|\bpresent\b)', duration.lower())
                    if year_match:
                        start_year = int(year_match.group(1)); end_year_str = year_match.group(2)
                        current_year = datetime.date.today().year 
                        end_year = current_year if end_year_str.lower() == "present" else int(end_year_str)
                        candidate_yoe += (end_year - start_year)
                    else:
                        direct_years = re.search(r'(\d+)\s*(?:year|yr|yrs|years)', duration.lower())
                        if direct_years: candidate_yoe += int(direct_years.group(1))
    context["candidate_yoe"] = candidate_yoe; context["jd_yoe_requirement"] = jd_data.get("required_experience", "N/A")
    num_companies = len(resume_data.get("experience", []))
    context["avg_tenure"] = round(candidate_yoe / num_companies, 1) if num_companies > 0 and candidate_yoe > 0 else "N/A"
    context["candidate_location"] = resume_data.get("location", "N/A"); context["jd_location"] = jd_data.get("location", "N/A")
    edu_summary = [f"{edu.get('degree','')} at {edu.get('institution','')}" for edu in resume_data.get("education", []) if isinstance(edu, dict)]
    context["candidate_education_summary"] = ", ".join(edu_summary) if edu_summary else "N/A"
    context["jd_education_summary"] = ", ".join(jd_data.get("education_requirements",[])) if jd_data.get("education_requirements") else jd_data.get("education_details", {}).get("degree_level", "N/A")
    cert_summary = [cert.get("name") for cert in resume_data.get("certifications", []) if isinstance(cert, dict) and cert.get("name")]
    context["candidate_certifications_summary"] = ", ".join(cert_summary) if cert_summary else "N/A"; context["jd_certifications_summary"] = "N/A"
    logger.debug(f"LLM evaluation context prepared: {json.dumps(context, indent=2)}") 
    return context

def create_llm_evaluation_prompt(resume_data_summary: Dict, jd_data_summary: Dict, pre_processed_context: Dict) -> str:
    prompt = f"""\
You are an expert Candidate-Job Fit Evaluator.
Analyze the provided Resume Summary and Job Description Summary, along with Pre-processed Context.
For each evaluation category, provide a numerical score (integer) within the specified maximum and a concise textual rationale.
Output your response as a single, valid JSON object.

Resume Summary (Key Information):
Name: {resume_data_summary.get("name", "N/A")}
Email: {resume_data_summary.get("email", "N/A")}
Phone: {resume_data_summary.get("phone", "N/A")}
Summary: {resume_data_summary.get("summary", "N/A")}
Total Experience: {pre_processed_context.get("candidate_yoe", "N/A")} years
Education: {pre_processed_context.get("candidate_education_summary", "N/A")}
Direct Skills: {list(resume_data_summary.get("direct_skills", {}).keys())}
Subjective Skills (examples): {list(resume_data_summary.get("subjective_skills", {}).keys())[:5]} 

Job Description Summary (Key Information):
Job Title: {jd_data_summary.get("job_title", "N/A")}
Company: {jd_data_summary.get("company_name", "N/A")}
Required Experience: {jd_data_summary.get("required_experience", "N/A")}
Required Skills: {jd_data_summary.get("required_skills", [])}
Preferred Skills: {jd_data_summary.get("preferred_skills", [])}
Education Requirements: {pre_processed_context.get("jd_education_summary", "N/A")}

Pre-processed Context:
- Matched Required Skills: {pre_processed_context.get("matched_required_skills", "N/A")}
- Missing Required Skills: {pre_processed_context.get("missing_required_skills", "N/A")}
- Matched Preferred Skills: {pre_processed_context.get("matched_preferred_skills", "N/A")}
- Subjective Skills (from experience/projects) relevant to JD: {pre_processed_context.get("subjective_skills_found_in_jd", "N/A")}
- Candidate's Average Job Tenure: {pre_processed_context.get("avg_tenure", "N/A")} years
- Candidate's Location: {pre_processed_context.get("candidate_location", "N/A")} vs JD Location: {pre_processed_context.get("jd_location", "N/A")}
- Candidate's Certifications: {pre_processed_context.get("candidate_certifications_summary", "N/A")}

Desired JSON Output Structure:
{{
  "direct_skills": {{ "score": <int>, "rationale": "<string>" }},
  "subjective_skills": {{ "score": <int>, "rationale": "<string>" }},
  "experience_match": {{ "score": <int>, "rationale": "<string>" }},
  "reliability": {{ "score": <int>, "rationale": "<string>" }},
  "location_match": {{ "score": <int>, "rationale": "<string>" }},
  "education_match": {{ "score": <int>, "rationale": "<string>" }},
  "certifications_match": {{ "score": <int>, "rationale": "<string>" }},
  "overall_summary": "<string>"
}}

Evaluation Categories & Max Scores:
1.  direct_skills: Match of candidate's explicitly listed skills with JD's required and preferred skills. (Max: 25)
2.  subjective_skills: Skills inferred from experience/projects matching JD, not covered by direct skills. (Max: 15)
3.  experience_match: Years of experience, relevance of roles, and quality compared to JD. (Max: 20)
4.  reliability: Job stability based on average tenure. (Max: 10)
5.  location_match: Candidate location vs. JD location. (Max: 10)
6.  education_match: Degree level, field, institution relevance to JD. (Max: 10)
7.  certifications_match: Relevance of certifications to JD. (Max: 5)
8.  overall_summary: Brief overall assessment (2-3 sentences).

Provide integer scores. Be critical and fair.
"""
    return prompt

async def evaluate_fit_with_llm(resume_data: Dict, jd_data: Dict, request_metrics: RequestMetrics = None) -> Dict:
    logger.info("Evaluating candidate-job fit with LLM")
    pre_processed_context = await asyncio.to_thread(_prepare_llm_evaluation_context, resume_data, jd_data)
    resume_summary_for_prompt = {"name": resume_data.get("name"), "email": resume_data.get("email"), "phone": resume_data.get("phone"), "summary": resume_data.get("summary"), "direct_skills": resume_data.get("direct_skills"), "subjective_skills": resume_data.get("subjective_skills")}
    jd_summary_for_prompt = {"job_title": jd_data.get("job_title"), "company_name": jd_data.get("company_name"), "required_experience": jd_data.get("required_experience"), "required_skills": jd_data.get("required_skills"), "preferred_skills": jd_data.get("preferred_skills")}
    prompt = create_llm_evaluation_prompt(resume_summary_for_prompt, jd_summary_for_prompt, pre_processed_context)
    try:
        response_text = await get_response(prompt, timeout_seconds=180, max_tokens=1024, request_metrics=request_metrics)
        json_str = response_text.strip()
        if "```json" in json_str: json_str = json_str.split("```json")[1].split("```")[0].strip()
        elif "```" in json_str: json_str = json_str.split("```")[1].split("```")[0].strip()
        if not (json_str.startswith('{') and json_str.endswith('}')): raise json.JSONDecodeError("LLM response is not a JSON object", json_str, 0)
        llm_evaluation = await asyncio.to_thread(json.loads, json_str)
        logger.info(f"LLM Evaluation Output: {llm_evaluation}")
        return llm_evaluation
    except json.JSONDecodeError as json_err:
        logger.error(f"LLM evaluation JSON decode error: {json_err.msg} for response: {response_text[:500] if 'response_text' in locals() else 'N/A'}", exc_info=True)
        return {"error": "LLM output parsing error", "details": str(json_err)}
    except Exception as e:
        logger.error(f"Error in LLM fit evaluation: {e}", exc_info=True) 
        return {"error": "LLM evaluation failed", "details": str(e)}

async def calculate_candidate_job_fit(resume_data: Dict, jd_data: Dict, request_metrics: RequestMetrics = None) -> Dict:
    logger.info("Calculating candidate-job fit score using LLM-based approach")
    llm_eval_results = await evaluate_fit_with_llm(resume_data, jd_data, request_metrics)
    scores = {}; rationale = {}; default_rationale = "LLM did not provide a specific rationale for this category."
    evaluation_categories_config = {
        "direct_skills": {"max_score": 25, "key_in_llm": "direct_skills"}, "subjective_skills": {"max_score": 15, "key_in_llm": "subjective_skills"},
        "experience_match": {"max_score": 20, "key_in_llm": "experience_match"}, "reliability": {"max_score": 10, "key_in_llm": "reliability"},
        "location_match": {"max_score": 10, "key_in_llm": "location_match"}, "education_match": {"max_score": 10, "key_in_llm": "education_match"},
        "certifications_match": {"max_score": 5, "key_in_llm": "certifications_match"}
    }
    if "error" in llm_eval_results:
        logger.error(f"LLM evaluation failed: {llm_eval_results.get('details')}")
        for cat_key, config in evaluation_categories_config.items():
            scores[cat_key] = 0; rationale[cat_key] = f"LLM evaluation error: {llm_eval_results.get('details', 'Unknown error')}"
        overall_summary_from_llm = "LLM evaluation failed."
    else:
        for internal_key, config in evaluation_categories_config.items():
            llm_category_data = llm_eval_results.get(config["key_in_llm"], {})
            score = llm_category_data.get("score")
            if isinstance(score, int) and 0 <= score <= config["max_score"]: scores[internal_key] = score
            else: logger.warning(f"Invalid or missing score for {internal_key} from LLM: {score}. Defaulting to 0."); scores[internal_key] = 0
            rationale[internal_key] = llm_category_data.get("rationale", default_rationale)
        overall_summary_from_llm = llm_eval_results.get("overall_summary", "LLM did not provide an overall summary.")
    total_score = sum(scores.values())
    fit_category = "Weak Match"
    if total_score >= 85: fit_category = "Excellent Match"
    elif total_score >= 70: fit_category = "Strong Match"
    elif total_score >= 55: fit_category = "Good Match"
    elif total_score >= 40: fit_category = "Moderate Match"
    summary = overall_summary_from_llm
    if summary == "LLM did not provide an overall summary." or "LLM evaluation failed" in summary :
         summary = f"The candidate is a {fit_category.lower()} for this position with a score of {total_score:.1f}/100."
         strengths = [cat.replace("_", " ").title() for cat, score in scores.items() if score >= evaluation_categories_config[cat]["max_score"] * 0.7]
         weaknesses = [cat.replace("_", " ").title() for cat, score in scores.items() if score <= evaluation_categories_config[cat]["max_score"] * 0.3]
         if strengths: summary += f" Key strengths: {', '.join(strengths)}."
         if weaknesses: summary += f" Areas for improvement: {', '.join(weaknesses)}."
    return {"total_score": total_score, "fit_category": fit_category, "summary": summary, "scores": scores, "rationale": rationale}

async def generate_jd_only_questions(jd_text: str) -> Dict:
    logger.info("Starting JD-only interview question generation (async)")
    categories = ["technical_questions", "past_experience_questions", "case_study_questions", "situation_handling_questions", "personality_test_questions"]
    questions_data = {}
    async def process_category(cat): return cat, await generate_questions_for_category(jd_text, None, cat, 5)
    results = await asyncio.gather(*(process_category(cat) for cat in categories))
    for category, questions in results: questions_data[category] = questions
    logger.info("Successfully generated all JD-only interview questions")
    return questions_data

@app.post("/jd_only", response_model=Dict)
async def generate_jd_only_questions_endpoint(request: Request, file: UploadFile = File(...)): 
    temp_file_path = None
    metrics = getattr(request.state, "metrics", None)
    try:
        file_extension = os.path.splitext(file.filename.lower())[1]
        if file_extension == '.pdf': file_type = "pdf"; suffix = '.pdf'
        elif file_extension == '.docx': file_type = "docx"; suffix = '.docx'
        else: raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
        content = await file.read()
        temp_file_path = await write_to_temp_file(content, suffix)
        jd_extraction_prompt = JD_PARSING_PROMPT_TEMPLATE.replace("... Job Description text...", "{text_content_placeholder}")
        jd_text = await extract_text_from_file(temp_file_path, file_type, jd_extraction_prompt, request_metrics=metrics)
        questions_data = await generate_jd_only_questions(jd_text)
        return questions_data
    except HTTPException as e: 
        if metrics: metrics.add_metric("error", str(e.detail if hasattr(e, 'detail') else e)); metrics.add_metric("error_status", e.status_code if hasattr(e, 'status_code') else 500)
        raise e
    except Exception as e: 
        if metrics: metrics.add_metric("error", str(e)); metrics.add_metric("error_status", 500)
        logger.error(f"Error in /jd_only for {file.filename}: {e}", exc_info=True); 
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred while processing {file.filename}: {str(e)}")
    finally:
        if temp_file_path and await asyncio.to_thread(os.path.exists, temp_file_path): await asyncio.to_thread(os.remove, temp_file_path)

@app.post("/intervet", response_model=Dict)
async def evaluate_candidate_job_fit_endpoint(request_body: IntervetRequest, request: Request):
    if not request_body.resume_json or not request_body.jd_json: raise HTTPException(status_code=400, detail="Missing or invalid JSON data.")
    try:
        metrics = getattr(request.state, "metrics", None)
        result = await calculate_candidate_job_fit(request_body.resume_json, request_body.jd_json, request_metrics=metrics)
        return result
    except Exception as e: 
        logger.error(f"Error in /intervet: {e}", exc_info=True); 
        if metrics: metrics.add_metric("error", str(e)); metrics.add_metric("error_status", 500)
        raise HTTPException(status_code=500, detail=f"An error occurred during candidate-job fit evaluation: {str(e)}")

@app.post("/bunchtest", response_model=Dict)
async def batch_evaluate_resumes_endpoint(request: Request, resume_files: List[UploadFile] = File(...), jd_file: UploadFile = File(...)):
    jd_temp_file_path = None; jd_data = {}
    metrics = getattr(request.state, "metrics", None)
    try:
        jd_file_extension = os.path.splitext(jd_file.filename.lower())[1]
        if jd_file_extension == '.pdf': jd_file_type = "pdf"; jd_suffix = '.pdf'
        elif jd_file_extension == '.docx': jd_file_type = "docx"; jd_suffix = '.docx'
        else: raise HTTPException(status_code=400, detail="JD file must be PDF or DOCX.")
        jd_content = await jd_file.read()
        jd_temp_file_path = await write_to_temp_file(jd_content, jd_suffix)
        
        specific_jd_parsing_prompt = JD_PARSING_PROMPT_TEMPLATE.replace("... Job Description text...", "{text_content_placeholder}")
        jd_text = await extract_text_from_file(jd_temp_file_path, jd_file_type, specific_jd_parsing_prompt, request_metrics=metrics)
        jd_data = await parse_jd_with_gemma(jd_text) 

        async def process_single_resume(resume_file_upload: UploadFile):
            loc_resume_temp_file = None
            try:
                r_file_extension = os.path.splitext(resume_file_upload.filename.lower())[1]
                if r_file_extension == '.pdf': loc_resume_file_type = "pdf"; r_suffix = '.pdf'
                elif r_file_extension == '.docx': loc_resume_file_type = "docx"; r_suffix = '.docx'
                else: logger.warning(f"Skipping {resume_file_upload.filename}: unsupported format."); return None
                r_content = await resume_file_upload.read()
                loc_resume_temp_file = await write_to_temp_file(r_content, r_suffix)
                
                specific_resume_parsing_prompt = RESUME_PARSING_PROMPT_TEMPLATE.replace("... Resume text...", "{text_content_placeholder}")
                loc_resume_text = await extract_text_from_file(loc_resume_temp_file, loc_resume_file_type, specific_resume_parsing_prompt, request_metrics=metrics)
                loc_resume_data = await parse_resume_with_gemma(loc_resume_text) 

                single_result = await calculate_candidate_job_fit(loc_resume_data, jd_data, request_metrics=metrics)
                single_result["resume_filename"] = resume_file_upload.filename
                return single_result
            except Exception as e_single: logger.error(f"Error processing resume {resume_file_upload.filename}: {e_single}", exc_info=True); return None 
            finally:
                if loc_resume_temp_file and await asyncio.to_thread(os.path.exists, loc_resume_temp_file): await asyncio.to_thread(os.remove, loc_resume_temp_file)
        
        results_gathered = await asyncio.gather(*(process_single_resume(rf) for rf in resume_files))
        results = [r for r in results_gathered if r]
        results.sort(key=lambda x: x.get("total_score", 0), reverse=True)
        return {"results": results, "jd_title": jd_data.get("job_title", "N/A"), "resume_count": len(results), "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")}
    except HTTPException as e_http: raise e_http
    except Exception as e_main: 
        logger.error(f"Error in /bunchtest: {e_main}", exc_info=True)
        if metrics: metrics.add_metric("error", str(e_main)); metrics.add_metric("error_status", 500)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during batch evaluation: {str(e_main)}")
    finally:
        if jd_temp_file_path and await asyncio.to_thread(os.path.exists, jd_temp_file_path): await asyncio.to_thread(os.remove, jd_temp_file_path)

@app.post("/interfix", response_model=InterfixResponse)
async def interfix_endpoint(request_body: InterfixRequest, request: Request):
    summary_text = request_body.summary
    metrics = getattr(request.state, "metrics", None)
    prompt = f"""You are an expert at extracting structured information... Call Summary/Transcript:\n{summary_text} ... Return ONLY a valid JSON object..."""
    try:
        response_text = await get_response(prompt, timeout_seconds=60, request_metrics=metrics)
        json_str = response_text.strip()
        if "```json" in json_str: json_str = json_str.split("```json")[1].split("```")[0].strip()
        elif "```" in json_str: json_str = json_str.split("```")[1].split("```")[0].strip()
        if not (json_str.startswith('{') and json_str.endswith('}')): raise json.JSONDecodeError("No JSON object found",json_str,0)
        extracted_data = await asyncio.to_thread(json.loads, json_str)
        return InterfixResponse(**extracted_data)
    except Exception as e: 
        logger.error(f"Error in /interfix: {e}", exc_info=True)
        if metrics: metrics.add_metric("error", str(e)); metrics.add_metric("error_status", 500)
        raise HTTPException(status_code=500, detail=f"An error occurred during interfix processing: {str(e)}")

@app.post("/intervet2", response_model=Dict)
async def evaluate_candidate_job_fit_from_files_endpoint(request: Request, resume_file: UploadFile = File(...), jd_file: UploadFile = File(...)):
    resume_temp_file_path = None; jd_temp_file_path = None
    metrics = getattr(request.state, "metrics", None)
    try:
        r_file_extension = os.path.splitext(resume_file.filename.lower())[1]
        if r_file_extension == '.pdf': resume_file_type = "pdf"; r_suffix = '.pdf'
        elif r_file_extension == '.docx': resume_file_type = "docx"; r_suffix = '.docx'
        else: raise HTTPException(status_code=400, detail="Resume file must be PDF or DOCX.")
        r_content = await resume_file.read()
        resume_temp_file_path = await write_to_temp_file(r_content, r_suffix)

        j_file_extension = os.path.splitext(jd_file.filename.lower())[1]
        if j_file_extension == '.pdf': jd_file_type = "pdf"; j_suffix = '.pdf'
        elif j_file_extension == '.docx': jd_file_type = "docx"; j_suffix = '.docx'
        else: raise HTTPException(status_code=400, detail="JD file must be PDF or DOCX.")
        j_content = await jd_file.read()
        jd_temp_file_path = await write_to_temp_file(j_content, j_suffix)
        
        specific_resume_parsing_prompt = RESUME_PARSING_PROMPT_TEMPLATE.replace("... Resume text...", "{text_content_placeholder}")
        resume_text = await extract_text_from_file(resume_temp_file_path, resume_file_type, specific_resume_parsing_prompt, request_metrics=metrics)
        resume_data_raw = await parse_resume_with_gemma(resume_text)
        resume_data = await normalize_resume_data(resume_data_raw, True)

        specific_jd_parsing_prompt = JD_PARSING_PROMPT_TEMPLATE.replace("... Job Description text...", "{text_content_placeholder}")
        jd_text = await extract_text_from_file(jd_temp_file_path, jd_file_type, specific_jd_parsing_prompt, request_metrics=metrics)
        jd_data_raw = await parse_jd_with_gemma(jd_text)
        jd_data = await normalize_jd_data(jd_data_raw)
        
        result = await calculate_candidate_job_fit(resume_data, jd_data, request_metrics=metrics)
        return result
    except HTTPException as e_http: 
        if metrics: metrics.add_metric("error", str(e_http.detail if hasattr(e_http, 'detail') else e_http)); metrics.add_metric("error_status", e_http.status_code if hasattr(e_http, 'status_code') else 500)
        raise e_http
    except Exception as e_main: 
        logger.error(f"Error in /intervet2 processing {resume_file.filename} and {jd_file.filename}: {e_main}", exc_info=True)
        if metrics: metrics.add_metric("error", str(e_main)); metrics.add_metric("error_status", 500)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e_main)}")
    finally:
        if resume_temp_file_path and await asyncio.to_thread(os.path.exists, resume_temp_file_path): await asyncio.to_thread(os.remove, resume_temp_file_path)
        if jd_temp_file_path and await asyncio.to_thread(os.path.exists, jd_temp_file_path): await asyncio.to_thread(os.remove, jd_temp_file_path)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

[end of main.py]

[end of main.py]
